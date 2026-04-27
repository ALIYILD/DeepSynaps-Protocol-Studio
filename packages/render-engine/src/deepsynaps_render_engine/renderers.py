"""Render engine — turn a ``ReportPayload`` into HTML, PDF, or DOCX.

Design notes
------------
* HTML is the source of truth — it is what the in-app viewer shows
  AND what the PDF wrapper passes through weasyprint.
* The HTML contains both the clinician view and the patient view in
  the same document; a small inline JS toggle flips visibility. Static
  PDF exports use a query-string flag (``?audience=patient``) handled
  by the calling router, but the default PDF includes both views
  separated by a page break.
* PDF requires ``weasyprint``. If not installed we raise a typed
  exception so the caller can return an HTTP 503 with a clear message
  rather than silently emitting a blank PDF.
* The pre-existing DOCX renderers (``render_protocol_docx``,
  ``render_patient_guide_docx``) remain unchanged for backwards
  compatibility with the worker.
"""

from __future__ import annotations

import html as _html
from io import BytesIO
from typing import Literal

from deepsynaps_core_schema import ProtocolPlan

from .payload import (
    CitationRef,
    InterpretationItem,
    ReportPayload,
    ReportSection,
    SuggestedAction,
)


# ---------------------------------------------------------------------------
# Errors
# ---------------------------------------------------------------------------


class RenderEngineError(RuntimeError):
    """Base exception for renderer failures."""


class PdfRendererUnavailable(RenderEngineError):
    """weasyprint (or its native deps) is not installed.

    Callers should map this to HTTP 503 with a clear message. We never
    return a blank PDF — see audit doc for the rationale.
    """


# ---------------------------------------------------------------------------
# Legacy exports (kept for backwards compatibility — used by IntakePreview)
# ---------------------------------------------------------------------------


def render_web_preview(protocol: ProtocolPlan) -> dict[str, object]:
    """Backwards-compatible shallow web preview for ProtocolPlan."""
    return {
        "title": protocol.title,
        "summary": protocol.summary,
        "checks": protocol.checks,
        "export_targets": ["web", "docx", "pdf"],
    }


# ---------------------------------------------------------------------------
# HTML renderer for ReportPayload
# ---------------------------------------------------------------------------


# Strength badge colours mirror the protocol-detail badge palette in
# pages-protocols.js so clinicians see one consistent visual language.
_STRENGTH_COLOR: dict[str, tuple[str, str]] = {
    # (text-color, bg-color)
    "Strong":           ("#0a5d2c", "#d1f7df"),
    "Moderate":         ("#9b6a00", "#fff2c8"),
    "Limited":          ("#7a3e00", "#fde2cc"),
    "Conflicting":      ("#7a1f1f", "#fbd5d5"),
    "Evidence pending": ("#475569", "#e2e8f0"),
}

_CONFIDENCE_LABEL: dict[str, str] = {
    "high":         "High confidence",
    "medium":       "Medium confidence",
    "low":          "Low confidence",
    "insufficient": "Insufficient evidence",
}


def _esc(text: object) -> str:
    return _html.escape("" if text is None else str(text), quote=True)


def _strength_badge(strength: str) -> str:
    color, bg = _STRENGTH_COLOR.get(strength, _STRENGTH_COLOR["Evidence pending"])
    return (
        f'<span class="ds-strength" '
        f'style="color:{color};background:{bg};padding:2px 8px;'
        f'border-radius:10px;font-size:11px;font-weight:600;'
        f'text-transform:uppercase;letter-spacing:0.3px;">{_esc(strength)}</span>'
    )


def _confidence_pill(level: str | None) -> str:
    if not level:
        return ""
    label = _CONFIDENCE_LABEL.get(level, level.title())
    color = {
        "high": "#0a5d2c",
        "medium": "#9b6a00",
        "low": "#7a3e00",
        "insufficient": "#475569",
    }.get(level, "#475569")
    return (
        f'<span class="ds-confidence" '
        f'style="color:{color};border:1px solid {color};padding:2px 10px;'
        f'border-radius:12px;font-size:11px;font-weight:600;">'
        f'{_esc(label)}</span>'
    )


def _citation_inline(refs: list[str], lookup: dict[str, CitationRef]) -> str:
    """Render `[C1, C3]`-style inline citation tags."""
    if not refs:
        return ""
    parts: list[str] = []
    for r in refs:
        cit = lookup.get(r)
        if cit is None:
            parts.append(f'<sup class="ds-cite ds-cite-missing">[{_esc(r)}?]</sup>')
            continue
        link = cit.best_link()
        body = f'[{_esc(r)}]'
        if link:
            parts.append(
                f'<sup class="ds-cite"><a href="{_esc(link)}" target="_blank" rel="noopener" '
                f'style="color:#1f5fb3;text-decoration:none;">{body}</a></sup>'
            )
        else:
            parts.append(f'<sup class="ds-cite ds-cite-unverified">{body}</sup>')
    return " " + "".join(parts)


def _render_observed(items: list[str]) -> str:
    if not items:
        return (
            '<div class="ds-empty" '
            'style="color:#64748b;font-style:italic;">No findings recorded for this section.</div>'
        )
    rows = "".join(f"<li>{_esc(i)}</li>" for i in items)
    return f'<ul class="ds-observed-list" style="margin:6px 0 0;padding-left:18px;">{rows}</ul>'


def _render_interpretations(
    items: list[InterpretationItem],
    lookup: dict[str, CitationRef],
) -> str:
    if not items:
        return (
            '<div class="ds-empty" '
            'style="color:#64748b;font-style:italic;">No model interpretations for this section.</div>'
        )
    out: list[str] = []
    for it in items:
        cites = _citation_inline(it.evidence_refs, lookup)
        counter = ""
        if it.counter_evidence_refs:
            counter = (
                ' <span class="ds-conflict" '
                'style="color:#7a1f1f;font-size:11px;font-weight:600;">'
                f'Conflicting evidence: {_esc(", ".join(it.counter_evidence_refs))}</span>'
            )
        out.append(
            '<li style="margin-bottom:8px;">'
            f'{_strength_badge(it.evidence_strength)} '
            f'<span>{_esc(it.text)}</span>{cites}{counter}</li>'
        )
    return f'<ul class="ds-interp-list" style="margin:6px 0 0;padding-left:18px;list-style:none;">{"".join(out)}</ul>'


def _render_suggestions(items: list[SuggestedAction]) -> str:
    if not items:
        return (
            '<div class="ds-empty" '
            'style="color:#64748b;font-style:italic;">No suggested actions.</div>'
        )
    out: list[str] = []
    for s in items:
        prefix = "Consider: " if s.requires_clinician_review else ""
        rationale = (
            f'<div class="ds-rationale" '
            f'style="color:#475569;font-size:12px;margin-top:2px;">'
            f'Why: {_esc(s.rationale)}</div>'
        ) if s.rationale else ""
        out.append(
            '<li style="margin-bottom:6px;">'
            f'<span>{_esc(prefix)}{_esc(s.text)}</span>{rationale}</li>'
        )
    return f'<ul class="ds-action-list" style="margin:6px 0 0;padding-left:18px;">{"".join(out)}</ul>'


def _render_simple_list(items: list[str], *, fallback: str, color: str) -> str:
    if not items:
        return f'<div class="ds-empty" style="color:{color};font-style:italic;">{fallback}</div>'
    return (
        f'<ul style="margin:4px 0 0;padding-left:18px;color:{color};">'
        + "".join(f"<li>{_esc(i)}</li>" for i in items)
        + "</ul>"
    )


def _render_section(
    section: ReportSection,
    lookup: dict[str, CitationRef],
) -> str:
    return (
        '<section class="ds-section" '
        'style="border:1px solid #e2e8f0;border-radius:10px;'
        'padding:14px 16px;margin-bottom:14px;background:#ffffff;">'
        '<header style="display:flex;align-items:center;justify-content:space-between;'
        'gap:10px;margin-bottom:10px;flex-wrap:wrap;">'
        f'<h3 style="margin:0;font-size:16px;color:#0f172a;">{_esc(section.title)}</h3>'
        f'{_confidence_pill(section.confidence)}'
        '</header>'
        # Observed
        '<div class="ds-block ds-block-observed" '
        'style="border-left:3px solid #1f5fb3;padding:6px 10px;background:#f4f9ff;'
        'border-radius:0 6px 6px 0;margin-bottom:10px;">'
        '<div class="ds-label" '
        'style="font-size:11px;font-weight:700;color:#1f5fb3;text-transform:uppercase;'
        'letter-spacing:0.4px;">Observed findings</div>'
        f'{_render_observed(section.observed)}'
        '</div>'
        # Interpretations
        '<div class="ds-block ds-block-interp" '
        'style="border-left:3px solid #9b6a00;padding:6px 10px;background:#fffaf0;'
        'border-radius:0 6px 6px 0;margin-bottom:10px;">'
        '<div class="ds-label" '
        'style="font-size:11px;font-weight:700;color:#9b6a00;text-transform:uppercase;'
        'letter-spacing:0.4px;">Model interpretation</div>'
        f'{_render_interpretations(section.interpretations, lookup)}'
        '</div>'
        # Suggestions
        '<div class="ds-block ds-block-action" '
        'style="border-left:3px solid #0a5d2c;padding:6px 10px;background:#f3fbf6;'
        'border-radius:0 6px 6px 0;margin-bottom:10px;">'
        '<div class="ds-label" '
        'style="font-size:11px;font-weight:700;color:#0a5d2c;text-transform:uppercase;'
        'letter-spacing:0.4px;">Suggested actions (decision support)</div>'
        f'{_render_suggestions(section.suggested_actions)}'
        '</div>'
        # Cautions / limitations — always rendered
        '<details open style="margin-top:8px;">'
        '<summary style="cursor:pointer;font-size:12px;color:#475569;font-weight:600;">'
        'Cautions, limitations &amp; conflicting evidence</summary>'
        '<div style="display:grid;grid-template-columns:1fr 1fr;gap:10px;margin-top:6px;">'
        '<div><div class="ds-label" '
        'style="font-size:11px;font-weight:700;color:#7a3e00;text-transform:uppercase;'
        'letter-spacing:0.4px;">Cautions</div>'
        f'{_render_simple_list(section.cautions, fallback="No cautions identified.", color="#7a3e00")}'
        '</div>'
        '<div><div class="ds-label" '
        'style="font-size:11px;font-weight:700;color:#7a1f1f;text-transform:uppercase;'
        'letter-spacing:0.4px;">Limitations</div>'
        f'{_render_simple_list(section.limitations, fallback="None recorded.", color="#7a1f1f")}'
        '</div>'
        '</div>'
        + (
            '<div style="margin-top:8px;"><div class="ds-label" '
            'style="font-size:11px;font-weight:700;color:#7a1f1f;text-transform:uppercase;'
            'letter-spacing:0.4px;">Conflicting evidence</div>'
            f'<div style="color:#7a1f1f;font-size:12px;">'
            f'{_esc(", ".join(section.counter_evidence_refs))}'
            '</div></div>'
            if section.counter_evidence_refs else ''
        )
        + '</details>'
        '</section>'
    )


def _render_citations(citations: list[CitationRef]) -> str:
    if not citations:
        return (
            '<div class="ds-empty" '
            'style="color:#64748b;font-style:italic;">No citations attached.</div>'
        )
    rows: list[str] = []
    for c in citations:
        link = c.best_link()
        link_html = (
            f'<a href="{_esc(link)}" target="_blank" rel="noopener" '
            f'style="color:#1f5fb3;text-decoration:none;">{_esc(link)}</a>'
        ) if link else (
            f'<span class="ds-cite-unverified" '
            f'style="color:#7a1f1f;font-style:italic;">unverified — {_esc(c.raw_text or "no link")}</span>'
        )
        status_pill = (
            '<span class="ds-cite-status" '
            f'style="font-size:10px;font-weight:700;text-transform:uppercase;'
            f'padding:1px 6px;border-radius:8px;'
            + (
                'color:#0a5d2c;background:#d1f7df;'
                if c.status == "verified" else
                'color:#7a1f1f;background:#fbd5d5;'
                if c.status == "retracted" else
                'color:#7a3e00;background:#fde2cc;'
            )
            + f'">{_esc(c.status)}</span>'
        )
        evidence_pill = (
            f'<span style="margin-left:6px;font-size:11px;color:#475569;">'
            f'evidence: {_esc(c.evidence_level)}</span>' if c.evidence_level else ''
        )
        retrieved = (
            f'<span style="margin-left:6px;font-size:11px;color:#64748b;">'
            f'retrieved {_esc(c.retrieved_at)}</span>' if c.retrieved_at else ''
        )
        authors = ", ".join(c.authors[:3]) + (" et al" if len(c.authors) > 3 else "")
        meta = " · ".join(
            x for x in [
                _esc(authors),
                _esc(c.year),
                _esc(c.journal),
            ] if x
        )
        rows.append(
            '<li style="margin-bottom:8px;border-bottom:1px solid #e2e8f0;padding-bottom:6px;">'
            f'<div><strong>[{_esc(c.citation_id)}]</strong> '
            f'{_esc(c.title or "(untitled)")} {status_pill}{evidence_pill}{retrieved}</div>'
            f'<div style="color:#64748b;font-size:12px;">{meta}</div>'
            f'<div style="font-size:12px;">{link_html}</div>'
            '</li>'
        )
    return (
        '<ol class="ds-citation-list" '
        'style="margin:6px 0 0;padding-left:18px;list-style:none;">'
        + "".join(rows) +
        '</ol>'
    )


def _render_view(payload: ReportPayload, *, audience: str) -> str:
    lookup = {c.citation_id: c for c in payload.citations}
    sections_html = "".join(_render_section(s, lookup) for s in payload.sections)
    cautions_html = _render_simple_list(
        payload.global_cautions,
        fallback="No global cautions identified.",
        color="#7a3e00",
    )
    limitations_html = _render_simple_list(
        payload.global_limitations,
        fallback="No global limitations recorded.",
        color="#7a1f1f",
    )
    citations_html = _render_citations(payload.citations)

    audience_label = (
        "Clinician view" if audience == "clinician"
        else "Patient view" if audience == "patient"
        else ""
    )

    return (
        f'<div class="ds-view ds-view-{_esc(audience)}" data-audience="{_esc(audience)}">'
        '<header style="margin-bottom:14px;">'
        f'<div style="font-size:11px;color:#64748b;font-weight:600;text-transform:uppercase;'
        f'letter-spacing:0.4px;">{_esc(audience_label)}</div>'
        f'<h1 style="margin:2px 0 6px;font-size:22px;color:#0f172a;">{_esc(payload.title)}</h1>'
        f'<p style="margin:0;color:#334155;line-height:1.5;">{_esc(payload.summary)}</p>'
        '</header>'
        f'{sections_html}'
        '<section class="ds-globals" '
        'style="border:1px solid #e2e8f0;border-radius:10px;padding:14px 16px;'
        'margin-bottom:14px;background:#fafbff;">'
        '<h3 style="margin:0 0 8px;font-size:14px;color:#0f172a;">Cautions &amp; limitations</h3>'
        '<div style="display:grid;grid-template-columns:1fr 1fr;gap:10px;">'
        f'<div><strong style="font-size:12px;color:#7a3e00;">Cautions</strong>{cautions_html}</div>'
        f'<div><strong style="font-size:12px;color:#7a1f1f;">Limitations</strong>{limitations_html}</div>'
        '</div></section>'
        '<section class="ds-citations" '
        'style="border:1px solid #e2e8f0;border-radius:10px;padding:14px 16px;'
        'margin-bottom:14px;background:#ffffff;">'
        '<h3 style="margin:0 0 8px;font-size:14px;color:#0f172a;">Citations</h3>'
        f'{citations_html}'
        '</section>'
        '<footer style="font-size:11px;color:#64748b;border-top:1px solid #e2e8f0;'
        'padding-top:8px;margin-top:8px;">'
        f'<div>{_esc(payload.decision_support_disclaimer)}</div>'
        f'<div style="margin-top:4px;font-family:ui-monospace,monospace;">'
        f'schema: {_esc(payload.schema_id)} · generator: {_esc(payload.generator_version)} · '
        f'generated: {_esc(payload.generated_at)}</div>'
        '</footer>'
        '</div>'
    )


def render_report_html(
    payload: ReportPayload,
    *,
    audience: Literal["clinician", "patient", "both"] | None = None,
) -> str:
    """Render the structured payload to a self-contained HTML document.

    Parameters
    ----------
    payload : ReportPayload
        The structured payload.
    audience : "clinician" | "patient" | "both" | None
        Override the payload's ``audience``. ``"both"`` produces both
        views in the same document with a small inline JS toggle —
        works in the in-app viewer; the PDF wrapper renders both views
        separated by a page break.

    Returns
    -------
    str
        HTML document, never empty.
    """
    target = audience or payload.audience
    if target not in ("clinician", "patient", "both"):
        target = "both"

    if target == "both":
        body = (
            '<div class="ds-toggle-bar" '
            'style="display:flex;gap:8px;justify-content:center;margin-bottom:14px;">'
            '<button type="button" data-ds-view="clinician" '
            'class="ds-toggle ds-toggle-active" '
            'style="padding:6px 14px;border-radius:18px;border:1px solid #1f5fb3;'
            'background:#1f5fb3;color:#ffffff;font-weight:600;cursor:pointer;">Clinician</button>'
            '<button type="button" data-ds-view="patient" '
            'class="ds-toggle" '
            'style="padding:6px 14px;border-radius:18px;border:1px solid #1f5fb3;'
            'background:#ffffff;color:#1f5fb3;font-weight:600;cursor:pointer;">Patient</button>'
            '</div>'
            f'{_render_view(payload, audience="clinician")}'
            f'<div style="page-break-before:always;"></div>'
            f'{_render_view(payload, audience="patient")}'
            '<script>'
            "(function(){var btns=document.querySelectorAll('.ds-toggle');"
            "var views=document.querySelectorAll('.ds-view');"
            "function show(name){views.forEach(function(v){"
            "v.style.display=v.getAttribute('data-audience')===name?'':'none';});"
            "btns.forEach(function(b){var on=b.getAttribute('data-ds-view')===name;"
            "b.classList.toggle('ds-toggle-active',on);"
            "b.style.background=on?'#1f5fb3':'#ffffff';"
            "b.style.color=on?'#ffffff':'#1f5fb3';});}"
            "btns.forEach(function(b){b.addEventListener('click',function(){"
            "show(b.getAttribute('data-ds-view'));});});show('clinician');"
            "})();"
            '</script>'
        )
    else:
        body = _render_view(payload, audience=target)

    return (
        '<!doctype html><html lang="en"><head><meta charset="utf-8">'
        f'<title>{_esc(payload.title)}</title>'
        '<style>'
        'body{font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,sans-serif;'
        'background:#f5f7fb;color:#0f172a;margin:0;padding:24px;line-height:1.4;}'
        '.ds-view{max-width:880px;margin:0 auto;}'
        '@media print{body{background:#ffffff;padding:0;}.ds-toggle-bar{display:none;}'
        '.ds-view{max-width:none;}}'
        '</style></head><body>'
        f'{body}'
        '</body></html>'
    )


# ---------------------------------------------------------------------------
# PDF wrapper
# ---------------------------------------------------------------------------


def render_report_pdf(payload: ReportPayload) -> bytes:
    """Render the payload to PDF bytes via weasyprint.

    Raises
    ------
    PdfRendererUnavailable
        If ``weasyprint`` (or its native deps) is not importable.
        Callers should map this to HTTP 503. We never silently produce
        a blank PDF — see audit doc.
    """
    try:
        from weasyprint import HTML  # type: ignore[import-not-found]
    except Exception as exc:  # pragma: no cover - exercised in tests via patching
        raise PdfRendererUnavailable(
            "PDF export requires the 'weasyprint' package (and Pango/Cairo "
            "system libs). Install it on the API host to enable PDF export."
        ) from exc

    html_str = render_report_html(payload, audience=payload.audience)
    buf = BytesIO()
    HTML(string=html_str).write_pdf(target=buf)
    out = buf.getvalue()
    if not out:
        # weasyprint returning empty bytes indicates a broken install — surface it.
        raise PdfRendererUnavailable("weasyprint produced an empty PDF — check install.")
    return out


# ---------------------------------------------------------------------------
# Legacy DOCX renderers (unchanged)
# ---------------------------------------------------------------------------


def render_protocol_docx(protocol_plan, handbook_plan=None) -> bytes:
    """
    Generate a DOCX document from a protocol plan.
    Returns raw bytes of the .docx file.
    Uses python-docx.
    """
    try:
        from docx import Document  # noqa: F401
        from docx.shared import Pt, RGBColor, Inches  # noqa: F401
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: F401
    except ImportError:
        raise ImportError("python-docx required: pip install python-docx")

    doc = Document()

    # Title
    title = doc.add_heading(level=0)
    title.text = f"Clinical Protocol: {getattr(protocol_plan, 'title', 'Protocol')}"

    # Summary section
    doc.add_heading("Protocol Summary", level=1)
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = "Table Grid"
    fields = [
        ("Condition", getattr(protocol_plan, 'condition_name', '')),
        ("Modality", getattr(protocol_plan, 'modality_name', '')),
        ("Device", getattr(protocol_plan, 'device_name', '')),
        ("Evidence Grade", getattr(protocol_plan, 'evidence_grade', '')),
    ]
    for i, (label, value) in enumerate(fields):
        row = summary_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value) if value else "N/A"

    doc.add_paragraph()

    # Approval badge / off-label notice
    approval = getattr(protocol_plan, 'approval_badge', None)
    if approval:
        p = doc.add_paragraph()
        p.add_run(f"Status: {approval}").bold = True

    # Contraindications
    contras = getattr(protocol_plan, 'contraindications', [])
    if contras:
        doc.add_heading("Contraindications", level=1)
        for item in contras:
            doc.add_paragraph(str(item), style="List Bullet")

    # Safety checks
    safety = getattr(protocol_plan, 'safety_checks', [])
    if safety:
        doc.add_heading("Safety Checks", level=1)
        for item in safety:
            doc.add_paragraph(str(item), style="List Bullet")

    # Session structure
    session = getattr(protocol_plan, 'session_structure', None)
    if session:
        doc.add_heading("Session Structure", level=1)
        steps = getattr(session, 'steps', [])
        for step in steps:
            step_title = getattr(step, 'title', '')
            step_desc = getattr(step, 'description', '')
            doc.add_heading(str(step_title), level=2)
            if step_desc:
                doc.add_paragraph(str(step_desc))

    # Handbook sections (if provided)
    if handbook_plan:
        sections = getattr(handbook_plan, 'sections', [])
        for section in sections:
            title_text = getattr(section, 'title', '')
            body = getattr(section, 'body', '')
            doc.add_heading(str(title_text), level=1)
            if body:
                doc.add_paragraph(str(body))

    # Disclaimer
    doc.add_heading("Disclaimer", level=1)
    doc.add_paragraph(
        "This document is a DRAFT support tool for qualified clinicians only. "
        "It does not constitute medical advice. All protocols require independent "
        "clinical review before application. DeepSynaps accepts no liability for "
        "clinical decisions made using this document."
    )

    # Save to bytes
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def render_patient_guide_docx(condition_name: str, modality_name: str, instructions: list[str]) -> bytes:
    """Generate a simple patient-facing guide."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx required")

    doc = Document()
    doc.add_heading(f"Your Treatment Guide: {condition_name}", level=0)
    doc.add_paragraph(f"Treatment approach: {modality_name}")
    doc.add_heading("What to Expect", level=1)
    for instruction in instructions:
        doc.add_paragraph(str(instruction), style="List Bullet")
    doc.add_paragraph(
        "Please discuss any questions or concerns with your clinician before, during, or after treatment."
    )
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


__all__ = [
    "RenderEngineError",
    "PdfRendererUnavailable",
    "render_web_preview",
    "render_report_html",
    "render_report_pdf",
    "render_protocol_docx",
    "render_patient_guide_docx",
]
