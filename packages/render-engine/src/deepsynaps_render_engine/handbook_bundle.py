"""Full clinician-review handbook bundle: DOCX, HTML, and PDF.

Content is assembled only from registry-backed ``HandbookDocument`` fields and
the companion ``ReportPayload`` — no invented parameters or citations.
"""

from __future__ import annotations

from io import BytesIO

from deepsynaps_core_schema import HandbookDocument

from .payload import ReportPayload
from .renderers import PdfRendererUnavailable, _render_view

HANDBOOK_AI_ASSISTED_DISCLAIMER = (
    "This AI-assisted handbook is a clinician-review draft. It supports protocol "
    "implementation planning only. It does not diagnose, prescribe, approve treatment, "
    "triage emergencies, or replace clinician judgement. All parameters and clinical use "
    "require clinician verification against local policy, device labelling, patient "
    "suitability, and current evidence."
)


def _doc_add_bullets(doc, items: list[str]) -> None:
    for item in items:
        if str(item).strip():
            doc.add_paragraph(str(item).strip(), style="List Bullet")


def render_handbook_bundle_docx(
    document: HandbookDocument,
    detailed_report: ReportPayload | None,
    *,
    condition_name: str,
    modality_name: str,
    device_name: str,
    handbook_kind_label: str,
    evidence_grade: str = "",
    approval_badge: str = "",
    generated_at: str | None = None,
) -> bytes:
    """Build a comprehensive protocol-application handbook as DOCX."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError:
        raise ImportError("python-docx required: pip install python-docx")

    doc = Document()

    # ── Cover ───────────────────────────────────────────────────────────────
    h0 = doc.add_heading("DeepSynaps Protocol Handbook", level=0)
    h0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "\n".join(
            [
                f"Condition: {condition_name}",
                f"Modality: {modality_name}",
                f"Device: {device_name or '—'}",
                f"Audience / handbook type: {handbook_kind_label}",
                f"Generated (UTC): {generated_at or '—'}",
                "Version: clinician-review draft (not a signed treatment order)",
            ]
        )
    )
    run.font.size = Pt(11)

    doc.add_heading("Safety disclaimer", level=1)
    doc.add_paragraph(HANDBOOK_AI_ASSISTED_DISCLAIMER)

    doc.add_heading("Protocol overview", level=1)
    doc.add_paragraph(document.title)
    doc.add_paragraph(document.overview)
    meta_tbl = doc.add_table(rows=5, cols=2)
    meta_tbl.style = "Table Grid"
    meta_rows = [
        ("Condition", condition_name),
        ("Modality", modality_name),
        ("Device", device_name or "—"),
        ("Evidence grade (registry)", evidence_grade or "Not graded in this export"),
        ("Registry / approval posture", approval_badge or "See clinician review"),
    ]
    for i, (k, v) in enumerate(meta_rows):
        meta_tbl.rows[i].cells[0].text = k
        meta_tbl.rows[i].cells[1].text = str(v)

    doc.add_heading("Eligibility and clinical framing", level=1)
    _doc_add_bullets(doc, list(document.eligibility))

    doc.add_heading("Setup checklist", level=1)
    _doc_add_bullets(doc, list(document.setup))

    doc.add_heading("Session workflow", level=1)
    _doc_add_bullets(doc, list(document.session_workflow))

    doc.add_heading("Parameters and operational notes", level=1)
    doc.add_paragraph(
        "Session timing, dosing, and device parameters appear only in the dataset-derived "
        "workflow lines above and in protocol registry exports. Do not apply stimulation "
        "parameters without verifying against device labelling, approved protocols, and "
        "patient-specific assessment."
    )

    doc.add_heading("Safety, contraindications, and adverse-event considerations", level=1)
    _doc_add_bullets(doc, list(document.safety))

    doc.add_heading("Monitoring and documentation", level=1)
    doc.add_paragraph(
        "Use clinic-standard scales, tolerability checks, and chart documentation. "
        "The structured report appendix lists suggested monitoring themes from the generator."
    )

    doc.add_heading("Troubleshooting", level=1)
    _doc_add_bullets(doc, list(document.troubleshooting))

    doc.add_heading("Escalation", level=1)
    _doc_add_bullets(doc, list(document.escalation))

    doc.add_heading("Patient-facing summary (review before sharing)", level=1)
    doc.add_paragraph(
        "Plain-language instructions must be adapted by the care team. The following "
        "summarises generator text only and is not a standalone patient instruction sheet."
    )
    doc.add_paragraph(document.overview)
    doc.add_paragraph(
        "Patients should contact the clinic for non-emergency questions; call emergency "
        "services for urgent or life-threatening symptoms."
    )

    doc.add_heading("References and source pointers", level=1)
    refs = [r for r in document.references if str(r).strip()]
    if refs:
        for r in refs:
            doc.add_paragraph(str(r).strip(), style="List Bullet")
    else:
        doc.add_paragraph("Evidence link unavailable in this environment.")

    doc.add_heading("Appendix: Detailed clinical report (structured)", level=1)
    doc.add_paragraph(
        "Structured ReportPayload sections separate observed findings, model interpretations, "
        "and suggested actions for clinician review."
    )
    if detailed_report is not None:
        lookup = {c.citation_id: c for c in detailed_report.citations}
        for sec in detailed_report.sections:
            doc.add_heading(sec.title, level=2)
            doc.add_paragraph("Observed findings:")
            _doc_add_bullets(doc, list(sec.observed))
            doc.add_paragraph("Model interpretation:")
            for it in sec.interpretations:
                doc.add_paragraph(f"[{it.evidence_strength}] {it.text}", style="List Bullet")
            doc.add_paragraph("Suggested actions (decision support):")
            for act in sec.suggested_actions:
                line = act.text
                if act.rationale:
                    line += f" — Rationale: {act.rationale}"
                doc.add_paragraph(line, style="List Bullet")
            if sec.cautions:
                doc.add_paragraph("Cautions:")
                _doc_add_bullets(doc, list(sec.cautions))
            if sec.limitations:
                doc.add_paragraph("Limitations:")
                _doc_add_bullets(doc, list(sec.limitations))
        if detailed_report.global_cautions:
            doc.add_paragraph("Global cautions:")
            _doc_add_bullets(doc, list(detailed_report.global_cautions))
        if detailed_report.global_limitations:
            doc.add_paragraph("Global limitations:")
            _doc_add_bullets(doc, list(detailed_report.global_limitations))
        doc.add_paragraph(detailed_report.decision_support_disclaimer or "")
    else:
        doc.add_paragraph("Structured report unavailable for this export.")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def render_handbook_bundle_html(
    document: HandbookDocument,
    detailed_report: ReportPayload | None,
    *,
    condition_name: str,
    modality_name: str,
    device_name: str,
    handbook_kind_label: str,
    evidence_grade: str = "",
    approval_badge: str = "",
    generated_at: str | None = None,
) -> str:
    """Self-contained HTML for PDF printing or preview."""
    cover = (
        '<div style="text-align:center;margin-bottom:28px;padding-bottom:18px;'
        'border-bottom:2px solid #1f5fb3;">'
        '<h1 style="margin:0;font-size:26px;color:#0f172a;">DeepSynaps Protocol Handbook</h1>'
        f'<p style="margin:12px 0 0;font-size:13px;color:#334155;line-height:1.6;">'
        f"<strong>Condition:</strong> {_html_esc(condition_name)}<br/>"
        f"<strong>Modality:</strong> {_html_esc(modality_name)}<br/>"
        f"<strong>Device:</strong> {_html_esc(device_name or '—')}<br/>"
        f"<strong>Handbook type:</strong> {_html_esc(handbook_kind_label)}<br/>"
        f"<strong>Generated (UTC):</strong> {_html_esc(generated_at or '—')}"
        "</p></div>"
    )

    disc = (
        '<section style="border:2px solid #b45309;background:#fffbeb;padding:14px 16px;'
        'margin-bottom:22px;border-radius:8px;">'
        '<h2 style="margin:0 0 8px;font-size:15px;color:#92400e;">Safety disclaimer</h2>'
        f'<p style="margin:0;color:#78350f;line-height:1.55;font-size:13px;">'
        f"{_html_esc(HANDBOOK_AI_ASSISTED_DISCLAIMER)}</p></section>"
    )

    overview = (
        '<section style="margin-bottom:20px;"><h2 style="color:#0f172a;">Protocol overview</h2>'
        f'<p style="color:#334155;line-height:1.55;"><strong>{_html_esc(document.title)}</strong></p>'
        f'<p style="color:#334155;line-height:1.55;">{_html_esc(document.overview)}</p>'
        '<table style="width:100%;border-collapse:collapse;font-size:13px;margin-top:10px;">'
        f"<tr><td style='border:1px solid #e2e8f0;padding:6px;'><strong>Evidence grade</strong></td>"
        f"<td style='border:1px solid #e2e8f0;padding:6px;'>{_html_esc(evidence_grade or '—')}</td></tr>"
        f"<tr><td style='border:1px solid #e2e8f0;padding:6px;'><strong>Approval posture</strong></td>"
        f"<td style='border:1px solid #e2e8f0;padding:6px;'>{_html_esc(approval_badge or '—')}</td></tr>"
        "</table></section>"
    )

    def block(title: str, lines: list[str]) -> str:
        lis = "".join(f"<li>{_html_esc(x)}</li>" for x in lines if str(x).strip())
        inner = lis or "<li><em>No rows returned — verify against source protocol.</em></li>"
        return (
            f'<section style="margin-bottom:18px;"><h2 style="color:#0f172a;">{_html_esc(title)}</h2>'
            f"<ul style=\"margin:6px 0 0;padding-left:20px;color:#334155;\">{inner}</ul></section>"
        )

    narrative = "".join(
        [
            block("Eligibility and clinical framing", list(document.eligibility)),
            block("Setup checklist", list(document.setup)),
            block("Session workflow", list(document.session_workflow)),
            block("Safety, contraindications, monitoring", list(document.safety)),
            block("Troubleshooting", list(document.troubleshooting)),
            block("Escalation", list(document.escalation)),
        ]
    )

    refs = [r for r in document.references if str(r).strip()]
    ref_items = []
    for r in refs:
        rs = str(r).strip()
        if rs.lower().startswith(("http://", "https://")):
            ref_items.append(
                f'<li style="word-break:break-all;"><a href="{_html_esc(rs)}" '
                f'target="_blank" rel="noopener noreferrer" style="color:#1f5fb3">{_html_esc(rs)}</a></li>'
            )
        else:
            ref_items.append(f"<li>{_html_esc(rs)}</li>")
    ref_sec = (
        '<section style="margin-bottom:18px;"><h2 style="color:#0f172a;">References</h2>'
        + (
            f'<ul style="margin:6px 0 0;padding-left:20px;">{"".join(ref_items)}</ul>'
            if ref_items
            else '<p style="color:#64748b;font-style:italic;">Evidence link unavailable in this environment.</p>'
        )
        + "</section>"
    )

    appendix = ""
    if detailed_report is not None:
        appendix = (
            '<div style="page-break-before:always;"></div>'
            '<h2 style="color:#0f172a;">Appendix: Detailed clinical report</h2>'
            + _render_view(detailed_report, audience="clinician")
        )
    else:
        appendix = (
            '<div style="page-break-before:always;"></div>'
            '<h2 style="color:#0f172a;">Appendix: Detailed clinical report</h2>'
            '<p style="color:#64748b;font-style:italic;">Structured report unavailable.</p>'
        )

    body = (
        '<!doctype html><html lang="en"><head><meta charset="utf-8">'
        '<meta name="viewport" content="width=device-width, initial-scale=1">'
        '<title>Protocol Handbook</title>'
        '<style>body{font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,sans-serif;'
        "background:#f8fafc;color:#0f172a;margin:0;padding:28px;line-height:1.45;}"
        "@media print{body{background:#fff;padding:12px;}}"
        "</style></head><body>"
        f"{cover}{disc}{overview}{narrative}{ref_sec}{appendix}"
        "</body></html>"
    )
    return body


def _html_esc(text: object) -> str:
    import html as _html

    return _html.escape("" if text is None else str(text), quote=True)


def render_handbook_bundle_pdf(
    document: HandbookDocument,
    detailed_report: ReportPayload | None,
    *,
    condition_name: str,
    modality_name: str,
    device_name: str,
    handbook_kind_label: str,
    evidence_grade: str = "",
    approval_badge: str = "",
    generated_at: str | None = None,
) -> bytes:
    """PDF via WeasyPrint; raises PdfRendererUnavailable when unavailable."""
    html_str = render_handbook_bundle_html(
        document,
        detailed_report,
        condition_name=condition_name,
        modality_name=modality_name,
        device_name=device_name,
        handbook_kind_label=handbook_kind_label,
        evidence_grade=evidence_grade,
        approval_badge=approval_badge,
        generated_at=generated_at,
    )
    try:
        from weasyprint import HTML  # type: ignore[import-not-found]
    except Exception as exc:  # pragma: no cover - env dependent
        raise PdfRendererUnavailable(
            "PDF export requires WeasyPrint and system libraries (Pango/Cairo). "
            "DOCX export remains available."
        ) from exc

    buf = BytesIO()
    HTML(string=html_str).write_pdf(target=buf)
    out = buf.getvalue()
    if not out:
        raise PdfRendererUnavailable("WeasyPrint produced an empty PDF.")
    return out


__all__ = [
    "HANDBOOK_AI_ASSISTED_DISCLAIMER",
    "render_handbook_bundle_docx",
    "render_handbook_bundle_html",
    "render_handbook_bundle_pdf",
]
