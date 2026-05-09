"""Tests for the DOCX renderers in deepsynaps_render_engine.renderers.

The DOCX paths use ``getattr(plan, '...', default)`` everywhere, so the
test fixtures don't need to be the real Pydantic ProtocolPlan — a
SimpleNamespace with the right attribute names is enough. This keeps
the tests fast and focused on the rendering branches.
"""

from __future__ import annotations

from io import BytesIO
from types import SimpleNamespace

import pytest

from deepsynaps_render_engine.renderers import (
    render_patient_guide_docx,
    render_protocol_docx,
)


def _open(blob: bytes):
    """Reload a generated DOCX and return the python-docx Document."""
    from docx import Document
    return Document(BytesIO(blob))


def _full_protocol():
    return SimpleNamespace(
        title="rTMS for MDD",
        condition_name="Major Depressive Disorder",
        modality_name="rTMS",
        device_name="MagVenture MagPro",
        evidence_grade="A",
        approval_badge="FDA-cleared (depression, treatment-resistant)",
        contraindications=["Active seizure disorder", "Metallic implant near coil"],
        safety_checks=["Confirm informed consent", "Verify hearing protection"],
        session_structure=SimpleNamespace(
            steps=[
                SimpleNamespace(title="Setup", description="Calibrate device."),
                SimpleNamespace(title="Treatment", description="Deliver rTMS."),
            ],
        ),
    )


def _handbook(sections: list | None = None):
    return SimpleNamespace(
        sections=sections or [
            SimpleNamespace(title="Eligibility", body="Adults 18+"),
            SimpleNamespace(title="Setup", body="Verify device calibration."),
        ],
    )


# ───────────────────────────── render_protocol_docx ────────────────────────


class TestRenderProtocolDocx:
    def test_returns_non_empty_bytes(self) -> None:
        blob = render_protocol_docx(_full_protocol())
        assert isinstance(blob, bytes)
        assert len(blob) > 0

    def test_returns_valid_docx(self) -> None:
        blob = render_protocol_docx(_full_protocol())
        doc = _open(blob)  # re-opening would raise if the bytes are invalid
        assert doc is not None

    def test_contains_protocol_title(self) -> None:
        blob = render_protocol_docx(_full_protocol())
        doc = _open(blob)
        text = "\n".join(p.text for p in doc.paragraphs)
        # Title heading is also a paragraph in python-docx.
        assert "rTMS for MDD" in text or any(
            "rTMS for MDD" in p.text for p in doc.paragraphs
        )

    def test_summary_table_has_four_rows(self) -> None:
        blob = render_protocol_docx(_full_protocol())
        doc = _open(blob)
        # First table is the summary table; pin to 4 rows × 2 cols.
        assert doc.tables
        summary = doc.tables[0]
        assert len(summary.rows) == 4
        labels = [r.cells[0].text for r in summary.rows]
        assert labels == ["Condition", "Modality", "Device", "Evidence Grade"]

    def test_missing_optional_fields_render_na(self) -> None:
        # Plan with only the title — every getattr() default kicks in.
        plan = SimpleNamespace(title="Bare protocol")
        blob = render_protocol_docx(plan)
        doc = _open(blob)
        # Summary table cells fall back to "N/A".
        cells = [r.cells[1].text for r in doc.tables[0].rows]
        assert all(c == "N/A" for c in cells)

    def test_contraindications_rendered_as_bullets(self) -> None:
        blob = render_protocol_docx(_full_protocol())
        doc = _open(blob)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Active seizure disorder" in text
        assert "Metallic implant near coil" in text

    def test_safety_checks_rendered(self) -> None:
        blob = render_protocol_docx(_full_protocol())
        doc = _open(blob)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Confirm informed consent" in text

    def test_session_steps_rendered(self) -> None:
        blob = render_protocol_docx(_full_protocol())
        doc = _open(blob)
        text = "\n".join(p.text for p in doc.paragraphs)
        # Step titles render as level-2 headings.
        assert any("Setup" in p.text for p in doc.paragraphs)
        assert any("Treatment" in p.text for p in doc.paragraphs)
        assert "Calibrate device." in text

    def test_handbook_plan_sections_appended(self) -> None:
        blob = render_protocol_docx(_full_protocol(), handbook_plan=_handbook())
        doc = _open(blob)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Eligibility" in text
        assert "Adults 18+" in text
        assert "Setup" in text
        assert "Verify device calibration." in text

    def test_no_handbook_plan_no_handbook_sections(self) -> None:
        blob = render_protocol_docx(_full_protocol(), handbook_plan=None)
        doc = _open(blob)
        text = "\n".join(p.text for p in doc.paragraphs)
        # Without handbook, "Adults 18+" body must NOT appear.
        assert "Adults 18+" not in text

    def test_approval_badge_rendered_when_present(self) -> None:
        blob = render_protocol_docx(_full_protocol())
        doc = _open(blob)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "FDA-cleared" in text

    def test_no_approval_badge_when_absent(self) -> None:
        plan = SimpleNamespace(title="P", approval_badge=None)
        blob = render_protocol_docx(plan)
        doc = _open(blob)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Status:" not in text

    def test_disclaimer_always_present(self) -> None:
        # Pin the load-bearing safety contract: the disclaimer is the
        # legal-risk shield that the worker depends on. It must appear
        # even on a bare-minimum plan.
        plan = SimpleNamespace(title="Bare plan")
        blob = render_protocol_docx(plan)
        doc = _open(blob)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "DRAFT support tool" in text
        assert "qualified clinicians" in text
        assert "DeepSynaps accepts no liability" in text

    def test_empty_session_steps_still_renders(self) -> None:
        plan = SimpleNamespace(
            title="P",
            session_structure=SimpleNamespace(steps=[]),
        )
        blob = render_protocol_docx(plan)
        # Re-opens cleanly even with no steps.
        _open(blob)


# ───────────────────────────── render_patient_guide_docx ───────────────────


class TestRenderPatientGuideDocx:
    def test_returns_non_empty_bytes(self) -> None:
        blob = render_patient_guide_docx(
            "Major Depressive Disorder", "rTMS", ["Eat lightly before sessions"],
        )
        assert isinstance(blob, bytes)
        assert len(blob) > 0

    def test_contains_condition_and_modality(self) -> None:
        blob = render_patient_guide_docx(
            "MDD", "rTMS", ["Stay hydrated"],
        )
        doc = _open(blob)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "MDD" in text
        assert "rTMS" in text

    def test_instructions_rendered_as_bullets(self) -> None:
        blob = render_patient_guide_docx(
            "MDD", "rTMS",
            ["Eat lightly", "Avoid alcohol 24h before"],
        )
        doc = _open(blob)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Eat lightly" in text
        assert "Avoid alcohol 24h before" in text

    def test_clinician_disclaimer_present(self) -> None:
        # Patient-facing guide must direct questions to the clinician —
        # pin the safety boundary.
        blob = render_patient_guide_docx("MDD", "rTMS", [])
        doc = _open(blob)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "discuss" in text.lower()
        assert "clinician" in text.lower()

    def test_empty_instructions_still_valid_docx(self) -> None:
        blob = render_patient_guide_docx("MDD", "rTMS", [])
        _open(blob)  # would raise if the file is malformed
