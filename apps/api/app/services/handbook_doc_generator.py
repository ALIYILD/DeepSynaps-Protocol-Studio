"""
DeepSynaps Protocol Studio -- Handbook Document Generation Service

Generates clinical handbooks as DOCX (python-docx), PDF (fpdf2),
Markdown, patient-friendly DOCX, and ZIP bundles.
All documents carry a mandatory safety disclaimer.

Usage:
    docx = await generate_handbook_docx(handbook)
    pdf  = await generate_handbook_pdf(handbook)
    md   = await generate_handbook_markdown(handbook)
    pt   = await generate_patient_guide_docx(handbook, "simple")
    zip  = await generate_handbook_bundle(handbook)
"""

from __future__ import annotations

import io, json, zipfile
from dataclasses import dataclass, field
from datetime import datetime

from app.utils.time_utils import utc_now
from enum import Enum
from typing import Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from fpdf import FPDF

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

MANDATORY_DISCLAIMER = (
    "Draft for clinician review. Educational decision-support only.\n"
    "Not a diagnosis, prescription, or emergency guidance.\n\n"
    "Generated on {date} by {author} at {clinic}.\n"
    "Evidence grade: {evidence_grade} | Provenance: {provenance}\n"
    "Version: {version} | Review state: {review_state}\n\n"
    "Clinician review and sign-off required before clinical use.\n"
    "Evidence should be verified against current literature."
)

CLINICIAN_SECTIONS = [
    "overview", "indications", "contraindications", "preparation",
    "session_workflow", "safety_checklist", "adverse_events",
    "evidence_appendix", "limitations", "sign_off"
]
PATIENT_SECTIONS = [
    "what_is_this", "what_to_expect", "preparation",
    "during_session", "after_session", "risks_benefits",
    "questions_to_ask", "contacts"
]
SOP_SECTIONS = [
    "purpose", "scope", "responsibilities", "procedure",
    "safety_protocols", "documentation", "quality_assurance",
    "review_schedule"
]

FONT_BODY, FONT_SIZE_BODY, FONT_SIZE_HEADING = "Calibri", Pt(11), Pt(16)
FONT_SIZE_SECTION, FONT_SIZE_PATIENT = Pt(14), Pt(13)
MARGIN_INCH, SECTION_COLOR = Inches(1), RGBColor(0x2E, 0x50, 0x90)

# ---------------------------------------------------------------------------
# Data models
# ---------------------------------------------------------------------------

class EvidenceGrade(str, Enum):
    GRADE_A, GRADE_B, GRADE_C, GRADE_D = "A", "B", "C", "D"

class ReviewState(str, Enum):
    DRAFT, UNDER_REVIEW, APPROVED, ARCHIVED = "draft", "under_review", "approved", "archived"

@dataclass
class EvidenceItem:
    """Single evidence citation linked to a handbook section."""
    citation: str = ""
    pmid: Optional[str] = None
    doi: Optional[str] = None
    grade: str = "C"
    summary: str = ""

@dataclass
class HandbookSection:
    """A section within a clinical handbook."""
    title: str = ""
    key: str = ""
    body: str = ""
    subsections: List["HandbookSection"] = field(default_factory=list)
    order: int = 0

@dataclass
class Handbook:
    """Clinical handbook data structure consumed by all generators."""
    title: str = "Untitled Handbook"
    clinic: str = "DeepSynaps Clinic"
    author: str = "DeepSynaps Protocol Studio"
    version: str = "1.0.0"
    date: str = ""
    review_state: str = "draft"
    evidence_grade: str = "C"
    provenance: str = "AI-assisted synthesis"
    sections: List[HandbookSection] = field(default_factory=list)
    evidence: List[EvidenceItem] = field(default_factory=list)
    patient_contacts: Dict[str, str] = field(default_factory=dict)

    def __post_init__(self):
        if not self.date:
            self.date = utc_now().strftime("%Y-%m-%d")

# ---------------------------------------------------------------------------
# Template Engine
# ---------------------------------------------------------------------------

class HandbookTemplate:
    """Template engine. Templates: default | minimal | comprehensive | patient | sop | research"""

    def __init__(self, template_name: str = "default"):
        self.name = template_name
        self.section_order = {
            "default": CLINICIAN_SECTIONS,
            "minimal": ["overview", "indications", "preparation", "session_workflow", "sign_off"],
            "comprehensive": CLINICIAN_SECTIONS, "patient": PATIENT_SECTIONS,
            "sop": SOP_SECTIONS, "research": CLINICIAN_SECTIONS,
        }.get(template_name, CLINICIAN_SECTIONS)

    def render_cover(self, hb: Handbook) -> str:
        return (f"# {hb.title}\n**{hb.clinic}**\n\n**Version:** {hb.version}  "
                f"\n**Date:** {hb.date}  \n**Author:** {hb.author}  "
                f"\n**Review State:** {hb.review_state}  \n**Evidence Grade:** {hb.evidence_grade}\n\n---\n\n")

    def render_section(self, sec: HandbookSection) -> str:
        lines = [f"## {sec.title}", "", sec.body, ""]
        for sub in sec.subsections:
            lines += [f"### {sub.title}", "", sub.body, ""]
        return "\n".join(lines)

    def render_evidence_appendix(self, evidence: List[EvidenceItem]) -> str:
        if not evidence:
            return "## Evidence Appendix\n\nNo evidence items recorded.\n"
        lines = ["## Evidence Appendix", ""]
        for i, e in enumerate(evidence, 1):
            lines += [f"{i}. **[{e.grade}]** {e.citation}"]
            if e.pmid: lines.append(f"   - PMID: {e.pmid}")
            if e.doi: lines.append(f"   - DOI: {e.doi}")
            if e.summary: lines.append(f"   - Summary: {e.summary}")
            lines.append("")
        return "\n".join(lines)

    def render_disclaimer(self, hb: Handbook) -> str:
        return MANDATORY_DISCLAIMER.format(
            date=hb.date, author=hb.author, clinic=hb.clinic,
            evidence_grade=hb.evidence_grade, provenance=hb.provenance,
            version=hb.version, review_state=hb.review_state)

    def render_metadata(self, hb: Handbook) -> str:
        return json.dumps({
            "generated_at": utc_now().isoformat(),
            "generator": "DeepSynaps Protocol Studio v1.0",
            "handbook_title": hb.title, "version": hb.version, "author": hb.author,
            "clinic": hb.clinic, "section_count": len(hb.sections),
            "evidence_count": len(hb.evidence), "review_state": hb.review_state,
            "evidence_grade": hb.evidence_grade, "template": self.name}, indent=2)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _safe(v: Optional[str], fallback: str = "") -> str:
    return v if v else fallback

def _shade(cell, hex_color: str) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    tc_pr.append(cell._element.makeelement(qn("w:shd"), {
        qn("w:fill"): hex_color, qn("w:val"): "clear"}))

def _style(run, size: Pt = None, bold: bool = False,
           color: RGBColor = None, name: str = FONT_BODY) -> None:
    run.font.name = name
    if size: run.font.size = size
    if bold: run.bold = True
    if color: run.font.color.rgb = color

def _add_disclaimer(doc: Document, hb: Handbook) -> None:
    doc.add_page_break()
    p = doc.add_paragraph()
    _style(p.add_run("SAFETY DISCLAIMER"), Pt(12), True, RGBColor(0xCC, 0x33, 0x00))
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    tmpl = HandbookTemplate()
    _style(p2.add_run(tmpl.render_disclaimer(hb)), Pt(9), False, RGBColor(0x66, 0x66, 0x66))
    p2.runs[0].font.italic = True

# ---------------------------------------------------------------------------
# 1. DOCX Generation
# ---------------------------------------------------------------------------

async def generate_handbook_docx(
    handbook: Handbook,
    template_name: str = "default",
    include_evidence: bool = True,
    include_disclaimer: bool = True,
) -> bytes:
    """Generate DOCX handbook: cover, TOC, sections, evidence table, disclaimer, metadata."""
    doc, tmpl = Document(), HandbookTemplate(template_name)
    for section in doc.sections:
        for attr in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
            setattr(section, attr, MARGIN_INCH)
        hp = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _style(hp.add_run(f"{handbook.clinic}  |  {handbook.title}"), Pt(8), False, RGBColor(0x80, 0x80, 0x80))
        fp = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run("Page ")
        r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        fld = r._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
        r._element.append(fld)

    # Cover
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style(p.add_run(handbook.title), FONT_SIZE_HEADING, True, SECTION_COLOR)
    doc.add_paragraph()
    for line in [handbook.clinic, f"Version: {handbook.version}   |   Date: {handbook.date}",
                 f"Author: {handbook.author}   |   Review State: {handbook.review_state}",
                 f"Evidence Grade: {handbook.evidence_grade}"]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _style(p.add_run(line), Pt(11), False, RGBColor(0x44, 0x44, 0x44))
    doc.add_page_break()

    # TOC
    p = doc.add_paragraph()
    _style(p.add_run("Table of Contents"), FONT_SIZE_SECTION, True, SECTION_COLOR)
    doc.add_paragraph("(Update field to refresh table of contents)", style="Intense Quote")
    doc.add_page_break()

    # Body
    key_map = {s.key: s for s in handbook.sections}
    for key in tmpl.section_order:
        sec = key_map.get(key)
        if not sec: continue
        p = doc.add_paragraph()
        _style(p.add_run(sec.title), FONT_SIZE_SECTION, True, SECTION_COLOR)
        if sec.body:
            for line in sec.body.split("\n"):
                if line.strip():
                    p = doc.add_paragraph(line.strip())
                    for r in p.runs: _style(r, FONT_SIZE_BODY)
        for sub in sec.subsections:
            p = doc.add_paragraph()
            _style(p.add_run(sub.title), Pt(12), True, RGBColor(0x33, 0x33, 0x33))
            if sub.body:
                for line in sub.body.split("\n"):
                    if line.strip():
                        p = doc.add_paragraph(line.strip())
                        for r in p.runs: _style(r, FONT_SIZE_BODY)
        doc.add_paragraph()

    # Evidence
    if include_evidence and handbook.evidence:
        doc.add_page_break()
        p = doc.add_paragraph()
        _style(p.add_run("Evidence Appendix"), FONT_SIZE_SECTION, True, SECTION_COLOR)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(["#", "Grade", "Citation", "Summary"]):
            cell = table.rows[0].cells[i]; cell.text = h
            for pr in cell.paragraphs:
                for r in pr.runs: _style(r, Pt(10), True, RGBColor(0xFF, 0xFF, 0xFF))
            _shade(cell, "2E5090")
        for idx, item in enumerate(handbook.evidence, 1):
            row = table.add_row().cells
            row[0].text, row[1].text, row[2].text, row[3].text = str(idx), item.grade, item.citation, item.summary
            for cell in row: _shade(cell, "F0F4F8" if idx % 2 == 0 else "FFFFFF")
        doc.add_paragraph()

    if include_disclaimer: _add_disclaimer(doc, handbook)

    # Metadata
    doc.add_paragraph()
    p = doc.add_paragraph()
    _style(p.add_run("Generation Metadata"), Pt(10), True, RGBColor(0x80, 0x80, 0x80))
    p = doc.add_paragraph()
    _style(p.add_run(tmpl.render_metadata(handbook)), Pt(8), False, RGBColor(0x80, 0x80, 0x80))

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()

# ---------------------------------------------------------------------------
# 2. PDF Generation (fpdf2)
# ---------------------------------------------------------------------------

class _HandbookPDF(FPDF):
    """Custom FPDF with clinic header, page-number footer, and styled chapters."""

    def __init__(self, hb: Handbook, **kwargs):
        super().__init__(**kwargs); self.hb = hb

    def header(self):
        self.set_font("Helvetica", "B", 9); self.set_text_color(128, 128, 128)
        self.cell(0, 8, f"{self.hb.clinic}  |  {self.hb.title}", align="R", ln=True); self.ln(2)

    def footer(self):
        self.set_y(-15); self.set_font("Helvetica", "I", 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")

    def ch_title(self, title: str):
        self.set_font("Helvetica", "B", 14); self.set_text_color(0x2E, 0x50, 0x90)
        self.cell(0, 10, title, ln=True); self.ln(2)

    def ch_body(self, body: str):
        self.set_font("Helvetica", "", 11); self.set_text_color(0x33, 0x33, 0x33)
        for line in body.split("\n"):
            if line.strip(): self.multi_cell(0, 6, line.strip())
        self.ln(4)

    def disclaimer(self):
        self.set_font("Helvetica", "B", 10); self.set_text_color(0xCC, 0x33, 0x00)
        self.cell(0, 8, "SAFETY DISCLAIMER", ln=True)
        self.set_font("Helvetica", "I", 9); self.set_text_color(0x66, 0x66, 0x66)
        self.multi_cell(0, 5, HandbookTemplate().render_disclaimer(self.hb)); self.ln(4)


async def generate_handbook_pdf(handbook: Handbook, include_evidence: bool = True) -> bytes:
    """Generate PDF with header/footer, sections, evidence table, disclaimer, metadata."""
    pdf = _HandbookPDF(handbook); pdf.add_page()

    # Cover
    pdf.set_font("Helvetica", "B", 20); pdf.set_text_color(0x2E, 0x50, 0x90)
    pdf.cell(0, 12, handbook.title, ln=True, align="C"); pdf.ln(6)
    pdf.set_font("Helvetica", "", 12); pdf.set_text_color(0x44, 0x44, 0x44)
    for line in [handbook.clinic, f"Version: {handbook.version}  |  Date: {handbook.date}",
                 f"Author: {handbook.author}  |  Review State: {handbook.review_state}",
                 f"Evidence Grade: {handbook.evidence_grade}"]:
        pdf.cell(0, 7, line, ln=True, align="C")
    pdf.ln(8)

    # Sections
    for sec in handbook.sections:
        if not sec.body and not sec.subsections: continue
        pdf.ch_title(sec.title)
        if sec.body: pdf.ch_body(sec.body)
        for sub in sec.subsections:
            pdf.set_font("Helvetica", "B", 12); pdf.set_text_color(0x33, 0x33, 0x33)
            pdf.cell(0, 7, sub.title, ln=True)
            if sub.body: pdf.ch_body(sub.body)

    # Evidence
    if include_evidence and handbook.evidence:
        pdf.add_page(); pdf.ch_title("Evidence Appendix")
        pdf.set_font("Helvetica", "B", 10); cw = [15, 20, 80, 75]
        for w, h in zip(cw, ["#", "Grade", "Citation", "Summary"]): pdf.cell(w, 7, h, border=1, align="C")
        pdf.ln(); pdf.set_font("Helvetica", "", 9)
        for idx, e in enumerate(handbook.evidence, 1):
            pdf.cell(cw[0], 6, str(idx), border=1); pdf.cell(cw[1], 6, e.grade, border=1)
            pdf.cell(cw[2], 6, e.citation[:45], border=1); pdf.cell(cw[3], 6, e.summary[:42], border=1); pdf.ln()
        pdf.ln(6)

    pdf.add_page(); pdf.disclaimer()
    pdf.set_font("Helvetica", "B", 9); pdf.set_text_color(128, 128, 128)
    pdf.cell(0, 6, "Generation Metadata", ln=True); pdf.set_font("Helvetica", "", 8)
    for k, v in {"generated_at": utc_now().isoformat(), "generator": "DeepSynaps Protocol Studio",
                 "section_count": len(handbook.sections), "evidence_count": len(handbook.evidence)}.items():
        pdf.cell(0, 5, f"{k}: {v}", ln=True)
    return pdf.output()

# ---------------------------------------------------------------------------
# 3. Markdown Export
# ---------------------------------------------------------------------------

async def generate_handbook_markdown(handbook: Handbook, include_evidence: bool = True) -> str:
    """Generate Markdown for Git/version control."""
    tmpl = HandbookTemplate(); lines = [tmpl.render_cover(handbook)]
    for sec in handbook.sections: lines.append(tmpl.render_section(sec))
    if include_evidence: lines.append(tmpl.render_evidence_appendix(handbook.evidence))
    lines += ["## Disclaimer", ""]
    for dl in tmpl.render_disclaimer(handbook).split("\n"): lines.append(f"> {dl}")
    lines += ["", "<!--", "GENERATION METADATA", tmpl.render_metadata(handbook), "-->"]
    return "\n".join(lines)

# ---------------------------------------------------------------------------
# 4. Patient-Friendly DOCX
# ---------------------------------------------------------------------------

_LVL = {
    "simple": {"font_size": Pt(14), "heading_size": Pt(18), "desc": "Simplified 6th-grade reading level"},
    "standard": {"font_size": Pt(13), "heading_size": Pt(16), "desc": "Standard 8th-grade reading level"},
    "advanced": {"font_size": Pt(12), "heading_size": Pt(14), "desc": "Advanced reading level"},
}

_INTROS = {
    "what_is_this": "This guide explains your treatment in simple words. It tells you what will happen and what to expect.",
    "what_to_expect": "Here is what will happen during your visit. We want you to feel comfortable and prepared.",
    "preparation": "Please follow these steps to get ready for your session. If you have questions, call us anytime.",
    "during_session": "During your session, our team will guide you every step of the way. You can ask questions at any time.",
    "after_session": "After your session, follow these simple steps to take care of yourself. Most people feel fine afterward.",
    "risks_benefits": "Every treatment has benefits and some possible risks. We will explain both so you can make an informed choice.",
    "questions_to_ask": "Here are good questions to ask your care team. Write down any other questions you have.",
    "contacts": "If you need help, use these contact numbers. In an emergency, call 911.",
}

_SUBS = {
    "contraindicated": "should not be used", "contraindication": "reason not to use",
    "administer": "give", "administration": "giving", "evaluate": "check", "assess": "check",
    "monitor": "watch", "adverse event": "side effect", "adverse effects": "side effects",
    "efficacy": "how well it works", "therapeutic": "treatment", "protocol": "plan",
    "intervention": "treatment", "concurrent": "at the same time", "concomitant": "at the same time",
    "precaution": "safety step", "prophylactic": "preventive", "contraindications": "reasons not to use",
    "indications": "reasons to use", "utilize": "use",
}

def _simplify(text: str, level: str) -> str:
    for term, repl in _SUBS.items(): text = text.replace(term, repl)
    if level == "simple": text = text.replace(".", ". ").replace("  ", " ")
    return text


async def generate_patient_guide_docx(
    handbook: Handbook, reading_level: str = "standard"
) -> bytes:
    """Generate simplified patient-friendly DOCX with larger font, more whitespace, FAQ format."""
    level = reading_level if reading_level in _LVL else "standard"
    cfg = _LVL[level]; doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1.25)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style(p.add_run(f"Your Guide: {handbook.title}"), cfg["heading_size"], True, SECTION_COLOR)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style(p.add_run(f"A patient-friendly guide from {handbook.clinic}"), Pt(12), False, RGBColor(0x55, 0x55, 0x55))
    p = doc.add_paragraph()
    _style(p.add_run(f"Reading level: {cfg['desc']}"), Pt(10), False, RGBColor(0x88, 0x88, 0x88))
    p.runs[0].font.italic = True
    doc.add_page_break()

    for sec in handbook.sections:
        key = sec.key or sec.title.lower().replace(" ", "_")
        p = doc.add_paragraph()
        _style(p.add_run(sec.title), cfg["heading_size"] - Pt(2), True, SECTION_COLOR)
        intro = _INTROS.get(key, "")
        if intro:
            p = doc.add_paragraph()
            _style(p.add_run(intro), cfg["font_size"], False, RGBColor(0x55, 0x55, 0x55))
            p.runs[0].font.italic = True
        if sec.body:
            for line in _simplify(sec.body, level).split("\n"):
                if line.strip():
                    p = doc.add_paragraph(line.strip())
                    for r in p.runs: _style(r, cfg["font_size"])
                    p.paragraph_format.space_after = Pt(10)
        for sub in sec.subsections:
            p = doc.add_paragraph()
            _style(p.add_run(f"Q: {sub.title}"), Pt(cfg["font_size"].pt + 1), True, RGBColor(0x33, 0x33, 0x33))
            if sub.body:
                p = doc.add_paragraph()
                _style(p.add_run(f"A: {_simplify(sub.body, level)}"), cfg["font_size"])
                p.paragraph_format.space_after = Pt(8)
        doc.add_paragraph()

    if handbook.patient_contacts:
        doc.add_page_break()
        p = doc.add_paragraph()
        _style(p.add_run("Contact Information"), cfg["heading_size"] - Pt(2), True, SECTION_COLOR)
        for label, value in handbook.patient_contacts.items():
            p = doc.add_paragraph()
            _style(p.add_run(f"{label}: "), cfg["font_size"], True)
            _style(p.add_run(value), cfg["font_size"])

    doc.add_page_break()
    p = doc.add_paragraph()
    _style(p.add_run("Important Note"), Pt(12), True, RGBColor(0xCC, 0x33, 0x00))
    p = doc.add_paragraph()
    _style(p.add_run(
        "This guide is for learning only. It does not replace advice from your doctor. "
        "Always talk to your care team about your health.\n\nIf you have an emergency, call 911 right away."),
        cfg["font_size"], False, RGBColor(0x66, 0x66, 0x66))

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()

# ---------------------------------------------------------------------------
# 5. Bundle Export (ZIP)
# ---------------------------------------------------------------------------

def _evidence_pdf(hb: Handbook) -> bytes:
    """Generate standalone evidence appendix PDF."""
    pdf = _HandbookPDF(hb); pdf.add_page(); pdf.ch_title("Evidence Appendix")
    for idx, e in enumerate(hb.evidence, 1):
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(0, 6, f"{idx}. [{e.grade}] {e.citation}", ln=True)
        pdf.set_font("Helvetica", "", 9)
        if e.pmid: pdf.cell(0, 5, f"   PMID: {e.pmid}", ln=True)
        if e.doi: pdf.cell(0, 5, f"   DOI: {e.doi}", ln=True)
        if e.summary: pdf.multi_cell(0, 5, f"   Summary: {e.summary}")
        pdf.ln(3)
    return pdf.output()


def _readme(hb: Handbook, formats: List[str], inc_ev: bool) -> str:
    """Generate README.txt for the ZIP bundle."""
    lines = ["=" * 60, "  DeepSynaps Protocol Studio -- Handbook Export Bundle", "=" * 60, "",
             f"Handbook:       {hb.title}", f"Clinic:         {hb.clinic}", f"Version:        {hb.version}",
             f"Date:           {hb.date}", f"Author:         {hb.author}", f"Review State:   {hb.review_state}",
             f"Evidence Grade: {hb.evidence_grade}", f"Sections:       {len(hb.sections)}",
             f"Evidence Items: {len(hb.evidence)}", f"Formats:        {', '.join(formats)}", "",
             "-" * 60, "FILES IN THIS BUNDLE", "-" * 60,
             "  handbook.docx       - Clinician handbook (DOCX)",
             "  handbook.pdf        - Clinician handbook (PDF)",
             "  handbook.md         - Version-control friendly Markdown",
             "  patient_guide.docx  - Patient-friendly simplified guide"]
    if inc_ev and hb.evidence:
        lines.append("  evidence_appendix.pdf - Standalone evidence citations")
    lines += ["  metadata.json       - Generation metadata (JSON)", "  README.txt          - This file", "",
              "-" * 60, "SAFETY DISCLAIMER", "-" * 60, "",
              HandbookTemplate().render_disclaimer(hb), "",
              "-" * 60, "Generated by DeepSynaps Protocol Studio v1.0", "=" * 60, ""]
    return "\n".join(lines)


async def generate_handbook_bundle(
    handbook: Handbook, formats: List[str] = None, include_evidence: bool = True
) -> bytes:
    """Generate ZIP bundle: docx + pdf + md + patient_guide + evidence + metadata.json + README.txt."""
    if formats is None: formats = ["docx", "pdf", "markdown"]
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        if "docx" in formats:
            zf.writestr("handbook.docx", await generate_handbook_docx(handbook, include_evidence=include_evidence))
        if "pdf" in formats:
            zf.writestr("handbook.pdf", await generate_handbook_pdf(handbook, include_evidence=include_evidence))
        if "markdown" in formats:
            zf.writestr("handbook.md", (await generate_handbook_markdown(handbook, include_evidence=include_evidence)).encode())
        zf.writestr("patient_guide.docx", await generate_patient_guide_docx(handbook))
        if include_evidence and handbook.evidence:
            zf.writestr("evidence_appendix.pdf", _evidence_pdf(handbook))
        zf.writestr("metadata.json", HandbookTemplate("default").render_metadata(handbook))
        zf.writestr("README.txt", _readme(handbook, formats, include_evidence))
    buf.seek(0)
    return buf.read()

# ---------------------------------------------------------------------------
# 6. Demo / Test
# ---------------------------------------------------------------------------

async def _demo() -> Dict[str, bytes]:
    """Generate sample documents with demo rTMS data. Returns dict of filename -> bytes."""
    hb = Handbook(
        title="rTMS for Treatment-Resistant Depression", clinic="NeuroWellness Clinic",
        author="Dr. A. Smith", version="2.1.0", review_state="under_review", evidence_grade="B",
        sections=[
            HandbookSection(key="overview", title="Overview", body=(
                "Repetitive Transcranial Magnetic Stimulation (rTMS) is a non-invasive "
                "neuromodulation technique for treatment-resistant depression.\n\n"
                "This protocol covers the standard 10-Hz left DLPFC protocol."),
                subsections=[HandbookSection(title="Mechanism of Action", body=(
                    "rTMS uses magnetic pulses to modulate neuronal activity in the "
                    "dorsolateral prefrontal cortex. High-frequency stimulation increases "
                    "cortical excitability."))]),
            HandbookSection(key="indications", title="Indications", body=(
                "- Major Depressive Disorder (single or recurrent episode)\n"
                "- Treatment-resistant (failure of >=2 antidepressant trials)\n"
                "- Ages 18-70\n- Hamilton Depression Rating Scale >= 20")),
            HandbookSection(key="contraindications", title="Contraindications", body=(
                "- Implanted metallic hardware in cranium\n- Cochlear implants\n"
                "- History of seizures or epilepsy\n- Active substance use disorder\n"
                "- Current pregnancy")),
            HandbookSection(key="session_workflow", title="Session Workflow", body=(
                "1. Motor threshold determination\n2. Coil positioning over left DLPFC\n"
                "3. 3000 pulses at 10 Hz, 120% MT\n4. Session duration: ~37 minutes\n"
                "5. Daily sessions, 5 days/week for 4-6 weeks")),
            HandbookSection(key="safety_checklist", title="Safety Checklist", body=(
                "Pre-session:\n- Verify no new implants or devices\n- Screen for substance use\n"
                "- Check seizure risk factors\n\nDuring session:\n"
                "- Monitor for headache or discomfort\n- Ensure emergency equipment is accessible\n\n"
                "Post-session:\n- Document any adverse events\n- Schedule next session")),
        ],
        evidence=[
            EvidenceItem(citation="O'Reardon et al. (2007) -- rTMS vs sham for MDD", pmid="17538199",
                         doi="10.1016/j.biopsych.2007.01.018", grade="A",
                         summary="Large RCT: rTMS superior to sham for MDD."),
            EvidenceItem(citation="George et al. (2010) -- rTMS durability study", pmid="20615403",
                         doi="10.1016/j.biopsych.2010.05.031", grade="B",
                         summary="Durability of rTMS response over 6 months."),
            EvidenceItem(citation="McClintock et al. (2018) -- rTMS network meta-analysis",
                         pmid="29146165", grade="A", summary="Network MA: rTMS ranked second for MDD after ECT."),
        ],
        patient_contacts={"Clinic Phone": "(555) 123-4567", "After-Hours": "(555) 987-6543",
                          "Email": "care@neurowellness.example"})
    return {
        "handbook.docx": await generate_handbook_docx(hb),
        "handbook.pdf": await generate_handbook_pdf(hb),
        "handbook.md": (await generate_handbook_markdown(hb)).encode(),
        "patient_guide.docx": await generate_patient_guide_docx(hb, "simple"),
        "bundle.zip": await generate_handbook_bundle(hb),
    }


if __name__ == "__main__":
    import asyncio
    async def main():
        total = 0
        for name, data in (await _demo()).items():
            print(f"  {name:22s} {len(data):>10,} bytes")
            total += len(data)
        print(f"  {'TOTAL':22s} {total:>10,} bytes")
    asyncio.run(main())
