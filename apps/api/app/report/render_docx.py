"""ReportDocument → Word .docx (python-docx)."""

from __future__ import annotations

import io
from typing import Any

from app.report.resolve import resolve_placeholders


def document_to_docx_bytes(doc: dict[str, Any], ctx: dict[str, Any]) -> bytes:
    import docx
    from docx.shared import Inches, Pt

    d = docx.Document()
    style = d.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = str(doc.get("title") or "EEG Report")
    d.add_heading(title, 0)

    for b in doc.get("blocks") or []:
        bt = b.get("type")
        if bt == "heading":
            lvl = min(2, max(0, int(b.get("level") or 1) - 1))
            text = resolve_placeholders(str(b.get("text") or ""), ctx)
            d.add_heading(text, lvl)
        elif bt == "paragraph":
            p = d.add_paragraph()
            _add_runs_resolved(p, str(b.get("text") or ""), ctx)
        elif bt == "patientCard":
            d.add_heading("Patient", 2)
            p = ctx.get("patient") or {}
            tbl = d.add_table(rows=4, cols=2)
            tbl.style = "Table Grid"
            rows = [
                ("First name", str(p.get("firstName") or "")),
                ("Last name", str(p.get("lastName") or "")),
                ("DOB", str(p.get("dob") or "")),
                ("Gender", str(p.get("gender") or "")),
            ]
            for i, (a, val) in enumerate(rows):
                tbl.rows[i].cells[0].text = a
                tbl.rows[i].cells[1].text = val
        elif bt == "findings":
            d.add_heading("EEG Findings", 2)
            p = d.add_paragraph()
            _add_runs_resolved(p, str(b.get("text") or ""), ctx)
        elif bt in ("spectraGrid", "erpFigure", "sourceFigure", "spikeSummary"):
            d.add_paragraph(f"[{bt} figure placeholder]", style="Intense Quote")
        elif bt == "indicesTable":
            d.add_heading("Indices", 2)
            nz = (ctx.get("indices") or {}).get("normativeZ")
            if isinstance(nz, dict) and nz:
                tbl = d.add_table(rows=1 + len(nz), cols=2)
                tbl.style = "Table Grid"
                hdr = tbl.rows[0].cells
                hdr[0].text = "Metric"
                hdr[1].text = "Value"
                for i, (k, v) in enumerate(nz.items()):
                    tbl.rows[i + 1].cells[0].text = str(k)
                    tbl.rows[i + 1].cells[1].text = str(v)
            else:
                d.add_paragraph("No normative indices on file.")
        elif bt == "conclusion":
            d.add_heading("Conclusion", 2)
            p = d.add_paragraph()
            _add_runs_resolved(p, str(b.get("text") or ""), ctx)
        elif bt == "recommendation":
            d.add_heading("Recommendation", 2)
            p = d.add_paragraph()
            _add_runs_resolved(p, str(b.get("text") or ""), ctx)
        elif bt == "signature":
            p = d.add_paragraph()
            _add_runs_resolved(p, str(b.get("text") or ""), ctx)
        elif bt == "pageBreak":
            d.add_page_break()
        elif bt == "figure":
            cap = resolve_placeholders(str(b.get("caption") or "Figure"), ctx)
            src = b.get("src")
            if src:
                try:
                    d.add_picture(io.BytesIO(_load_image_bytes(str(src))), width=Inches(5))
                except Exception:
                    d.add_paragraph(f"[Image unavailable: {cap}]")
            else:
                d.add_paragraph(f"[{cap}]")
        else:
            d.add_paragraph(str(b))

    bio = io.BytesIO()
    d.save(bio)
    return bio.getvalue()


def _add_runs_resolved(paragraph: Any, text: str, ctx: dict[str, Any]) -> None:
    plain = resolve_placeholders(text, ctx)
    paragraph.add_run(_strip_html(plain))


def _strip_html(s: str) -> str:
    import re

    s = re.sub(r"<[^>]+>", "", s)
    return s.replace("&nbsp;", " ")


def _load_image_bytes(src: str) -> bytes:
    """HTTP(S) URL or data URL — minimal."""
    if src.startswith("data:image"):
        import base64

        _, b64 = src.split(",", 1)
        return base64.b64decode(b64)
    raise ValueError("only embedded data URLs supported in MVP")
